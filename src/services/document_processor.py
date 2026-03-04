"""Document Processor - PDF and DOCX Analysis"""
import os
from pathlib import Path
from typing import Optional, Dict, Any
from src.utils.logging import logger

class DocumentProcessor:
    """Process PDF and DOCX documents to extract text"""
    
    def __init__(self):
        self.pdf_available = False
        self.docx_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required libraries are available"""
        try:
            import pdfplumber
            self.pdf_available = True
            logger.info("pdfplumber available for PDF processing")
        except ImportError:
            logger.warning("pdfplumber not installed. Run: pip install pdfplumber")
        
        try:
            import docx
            self.docx_available = True
            logger.info("python-docx available for DOCX processing")
        except ImportError:
            logger.warning("python-docx not installed. Run: pip install python-docx")
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dict with extracted text and metadata
        """
        if not self.pdf_available:
            return {
                "success": False,
                "error": "pdfplumber not available. Install with: pip install pdfplumber",
                "text": None
            }
        
        try:
            import pdfplumber
            
            logger.info(f"Processing PDF file: {file_path}")
            
            text_content = []
            page_count = 0
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                    
                    # Log progress for large PDFs
                    if page_num % 10 == 0:
                        logger.info(f"Processed {page_num}/{page_count} pages")
            
            full_text = "\n\n".join(text_content)
            
            logger.info(f"PDF processing complete. Pages: {page_count}, Characters: {len(full_text)}")
            
            return {
                "success": True,
                "text": full_text,
                "page_count": page_count,
                "file_type": "pdf",
                "char_count": len(full_text)
            }
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": None
            }
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with extracted text and metadata
        """
        if not self.docx_available:
            return {
                "success": False,
                "error": "python-docx not available. Install with: pip install python-docx",
                "text": None
            }
        
        try:
            import docx
            
            logger.info(f"Processing DOCX file: {file_path}")
            
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_text:
                full_text += "\n\n=== Tables ===\n\n" + "\n".join(table_text)
            
            paragraph_count = len(paragraphs)
            table_count = len(doc.tables)
            
            logger.info(f"DOCX processing complete. Paragraphs: {paragraph_count}, Tables: {table_count}, Characters: {len(full_text)}")
            
            return {
                "success": True,
                "text": full_text,
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "file_type": "docx",
                "char_count": len(full_text)
            }
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": None
            }
    
    def process_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Read plain text file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Dict with text content and metadata
        """
        try:
            logger.info(f"Processing TXT file: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            line_count = text.count('\n') + 1
            
            logger.info(f"TXT processing complete. Lines: {line_count}, Characters: {len(text)}")
            
            return {
                "success": True,
                "text": text,
                "line_count": line_count,
                "file_type": "txt",
                "char_count": len(text)
            }
            
        except Exception as e:
            logger.error(f"TXT processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "text": None
            }
    
    def process_document(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """
        Process document based on file extension
        
        Args:
            file_path: Path to document file
            file_extension: File extension (pdf, docx, txt)
            
        Returns:
            Dict with extracted text and metadata
        """
        file_extension = file_extension.lower().replace('.', '')
        
        if file_extension == 'pdf':
            return self.process_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return self.process_docx(file_path)
        elif file_extension == 'txt':
            return self.process_txt(file_path)
        else:
            return {
                "success": False,
                "error": f"Unsupported document type: {file_extension}",
                "text": None
            }


# Global instance
document_processor = DocumentProcessor()
